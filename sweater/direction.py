from docx import Document
from docx.shared import Pt

# Краткая хар-ка про СП
unit = '4'
departament = 'Энергетическое отделение'
director_of_sp = 'А.Н. Ниматов'


def add_text_to_paragraph(paragraph, text, bold=False, underline=False, size=12):
    run = paragraph.add_run(text)
    run.bold = bold
    run.underline = underline
    run.font.size = Pt(size)


def direction_completion(row):
    # Открываем документ
    doc = Document('docs_templates/direction_template.docx')

    add_text_to_paragraph(doc.paragraphs[7], unit, size=12)
    add_text_to_paragraph(doc.paragraphs[8], f'({departament})', size=12)
    add_text_to_paragraph(doc.paragraphs[11], str(row[0]), bold=True, size=14)
    add_text_to_paragraph(doc.paragraphs[13], row[1], underline=True, size=14)
    add_text_to_paragraph(doc.paragraphs[14], row[2], underline=True, size=14)
    add_text_to_paragraph(doc.paragraphs[17], unit, underline=True)
    add_text_to_paragraph(doc.paragraphs[21], row[5], underline=True)
    add_text_to_paragraph(doc.paragraphs[24], row[3], underline=True)
    add_text_to_paragraph(doc.paragraphs[27], row[7], underline=True)
    date = row[6].split('-')
    text = f'Сроки практики с «{date[0]}» по «{date[1]}»'
    add_text_to_paragraph(doc.paragraphs[29], text)
    add_text_to_paragraph(doc.paragraphs[32], row[8])
    add_text_to_paragraph(doc.paragraphs[36], director_of_sp)

    full_name = row[1].split()
    production_practice = row[5].split()
    data_to_save = f'{full_name[0]} {full_name[1][0]}{full_name[2][0]} напр {production_practice[0]}'
    doc.save(f'students/Направления/{data_to_save}.docx')