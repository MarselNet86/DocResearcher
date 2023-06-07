from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import time


def practice_completion(practice_period, organization, director, specialization, phone_number, fax_number, site, inn, kpp, orgn, description, stock_full_name, type_of_practice, address):
    doc = Document('docs_templates/contract_template.docx')
    par_0 = doc.paragraphs[1]
    run_0 = par_0.add_run('4')
    run_0.bold = True
    run_0.font.size = Pt(9)

    par_1 = doc.paragraphs[5]
    run_1 = par_1.add_run('г. Сургут')
    run_1.font.size = Pt(10)

    par_2 = doc.paragraphs[5]
    run_2 = par_2.add_run(' ' * 136)
    run_2.font.size = Pt(10)

    par_3 = doc.paragraphs[5]
    run_3 = par_3.add_run('«')
    run_3.font.size = Pt(10)

    date = practice_period.split('-')

    par_4 = doc.paragraphs[5]
    run_4 = par_4.add_run(date[0])
    run_4.underline = True
    run_4.font.size = Pt(10)

    par_5 = doc.paragraphs[5]
    run_5 = par_5.add_run('»  ')
    run_5.font.size = Pt(10)

    par_6 = doc.paragraphs[5]
    run_6 = par_6.add_run(date[1])
    run_6.underline = True
    run_6.font.size = Pt(10)

    par_7 = doc.paragraphs[5]
    run_7 = par_7.add_run('  2023 г.')
    run_7.font.size = Pt(10)
    text = '           Автономное учреждение профессионального образования Ханты-Мансийского автономного округа - Югры «Сургутский политехнический колледж», в лице  руководителя структурного подразделения (заместителя директора по УР) Ниматова Ахмеда Нажмудтиновича,  действующего на основании Доверенности от 10 июня 2022 г. № 3/22, именуемый в дальнейшем «Колледж», с одной стороны, и '

    run_8 = doc.paragraphs[8]
    run_8 = run_8.add_run(text)
    run_8.font.size = Pt(11)

    run_9 = doc.paragraphs[8]
    run_9 = run_9.add_run(organization)
    run_9.underline = True
    run_9.font.size = Pt(11)

    run_10 = doc.paragraphs[8]
    run_10 = run_10.add_run(' именуемое в дальнейшем «Профильная организация» в лице ')
    run_10.font.size = Pt(11)

    run_11 = doc.paragraphs[8]
    run_11 = run_11.add_run(director)
    run_11.underline = True
    run_11.font.size = Pt(11)

    run_12 = doc.paragraphs[8]
    run_12 = run_12.add_run(',  действующего на основании Устава, с другой стороны, заключили настоящий договор о нижеследующем:')
    run_12.font.size = Pt(11)
    table = doc.tables[0]
    table.cell(1, 1).text = f'{organization}' \
                            f'\nАдрес: {address}' \
                            f'\nТелефон: {phone_number}' \
                            f'\nФакс: {fax_number}' \
                            f'\nСайт: {site}' \
                            f'\nИНН: {inn}' \
                            f'\nКПП: {kpp}' \
                            f'\nОРГН: {orgn}' \
                            f'\n{description}'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    table = doc.tables[1]
    paragraph = table.cell(2, 0).add_paragraph(specialization)
    font = paragraph.runs[0].font
    font.size = Pt(11)
    paragraph = table.cell(2, 1).add_paragraph(specialization)
    font = paragraph.runs[0].font
    font.size = Pt(11)
    paragraph = table.cell(2, 2).add_paragraph(f'1 \n\n({stock_full_name})')
    font = paragraph.runs[0].font
    font.size = Pt(11)
    paragraph = table.cell(2, 3).add_paragraph(date[0])
    font = paragraph.runs[0].font
    font.size = Pt(11)
    paragraph = table.cell(2, 4).add_paragraph(date[1])
    font = paragraph.runs[0].font
    font.size = Pt(11)

    table = doc.tables[2]
    cell = table.cell(1, 0)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.text = organization
    table.cell(3, 0).text = address
    table.cell(4, 0).text = 'Генеральный директор'
    cell = table.cell(5, 0)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    initials = director.split()
    paragraph.text = f'{initials[0][0]}. {initials[1][0]}. {initials[2]}'
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    table = doc.tables[3]
    paragraph = table.cell(1, 0).add_paragraph(type_of_practice)
    font = paragraph.runs[0].font
    font.size = Pt(11)
    paragraph = table.cell(1, 1).add_paragraph('Административное здание')
    font = paragraph.runs[0].font
    font.size = Pt(11)
    paragraph = table.cell(1, 2).add_paragraph(address)
    font = paragraph.runs[0].font
    font.size = Pt(11)
    paragraph = table.cell(1, 3).add_paragraph(address)
    font = paragraph.runs[0].font
    font.size = Pt(11)

    table = doc.tables[4]
    cell = table.cell(1, 0)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.text = organization
    table.cell(3, 0).text = address
    table.cell(4, 0).text = 'Генеральный директор'
    cell = table.cell(5, 0)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.text = f'{initials[0][0]}. {initials[1][0]}. {initials[2]}'

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
    full_name = stock_full_name.split()
    data_to_save = f'{full_name[0]} {full_name[1][0]}{full_name[2][0]} ПП'
    doc.save(f'students/ПП/{data_to_save}.docx')